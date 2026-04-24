from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Inches
from supabase import create_client, Client
import pandas as pd
import io
import fitz
import os
import requests
from typing import Optional, List, Dict, Any, Tuple
from datetime import datetime, timezone
import uuid
import time
import json
import base64
import mimetypes
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

load_dotenv()

app = FastAPI(title="AI Document Platform API")

FRONTEND_URL = os.getenv("FRONTEND_URL", "https://ai-doc-platform-zeta.vercel.app")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_SERVICE_ROLE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
STORAGE_BUCKET = os.getenv("SUPABASE_STORAGE_BUCKET", "uploads")
ML_API_URL = os.getenv("ML_API_URL")

if not OPENAI_API_KEY:
    raise ValueError("Missing OPENAI_API_KEY")

if not SUPABASE_URL or not SUPABASE_SERVICE_ROLE_KEY:
    raise ValueError("Missing Supabase environment variables")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        FRONTEND_URL,
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

client = OpenAI(api_key=OPENAI_API_KEY)
supabase: Client = create_client(SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY)


class AskRequest(BaseModel):
    question: str
    user_id: str
    conversation_id: Optional[str] = None
    filename: Optional[str] = None


class CreateConversationRequest(BaseModel):
    user_id: str
    title: Optional[str] = "New Chat"
    filename: Optional[str] = None


class UpdateConversationRequest(BaseModel):
    title: Optional[str] = None
    filename: Optional[str] = None


class ExportReportRequest(BaseModel):
    filename: str
    user_id: str


class ExportDocxWithSnapshotRequest(BaseModel):
    filename: str
    user_id: str
    conversation_id: Optional[str] = None
    snapshot_base64: Optional[str] = None


class LoanApplication(BaseModel):
    Age: float
    Income: float
    LoanAmount: float
    CreditScore: float
    MonthsEmployed: float
    NumCreditLines: float
    InterestRate: float
    LoanTerm: float
    DTIRatio: float
    Education: str
    EmploymentType: str
    MaritalStatus: str
    HasMortgage: str
    HasDependents: str
    LoanPurpose: str
    HasCoSigner: str


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def chunk_text(text: str, size: int = 1000) -> List[str]:
    return [text[i:i + size] for i in range(0, len(text), size) if text[i:i + size].strip()]


def get_embedding(text: str) -> List[float]:
    res = client.embeddings.create(
        model="text-embedding-3-small",
        input=text
    )
    return res.data[0].embedding


def safe_json_value(value: Any) -> Any:
    try:
        if pd.isna(value):
            return None
    except Exception:
        pass

    if isinstance(value, pd.Timestamp):
        return value.isoformat()

    if isinstance(value, datetime):
        return value.isoformat()

    if hasattr(value, "item"):
        try:
            return value.item()
        except Exception:
            return str(value)

    return value


def clean_preview_df(df: pd.DataFrame, max_rows: int = 25, max_cols: int = 12) -> pd.DataFrame:
    preview = df.copy()
    preview = preview.iloc[:max_rows, :max_cols]
    preview.columns = [str(col) for col in preview.columns]

    for col in preview.columns:
        try:
            if pd.api.types.is_datetime64_any_dtype(preview[col]):
                preview[col] = preview[col].astype(str)
        except Exception:
            pass

    return preview


def dataframe_to_preview_rows(df: pd.DataFrame, max_rows: int = 25, max_cols: int = 12) -> Dict[str, Any]:
    preview = clean_preview_df(df, max_rows=max_rows, max_cols=max_cols)
    rows = []

    for _, row in preview.iterrows():
        item = {}
        for col in preview.columns:
            item[str(col)] = safe_json_value(row[col])
        rows.append(item)

    return {
        "columns": [str(c) for c in preview.columns],
        "rows": rows
    }


def detect_interesting_numeric_columns(df: pd.DataFrame) -> List[str]:
    numeric_cols = df.select_dtypes(include="number").columns.tolist()
    preferred_keywords = [
        "amount", "total", "price", "cost", "sales", "revenue",
        "balance", "qty", "quantity", "count", "score", "loan", "income"
    ]

    ranked = []
    for col in numeric_cols:
        lower = str(col).lower()
        score = 0
        for keyword in preferred_keywords:
            if keyword in lower:
                score += 5
        score += max(0, len(df[col].dropna()))
        ranked.append((col, score))

    ranked.sort(key=lambda x: x[1], reverse=True)
    return [col for col, _ in ranked[:6]]


def infer_date_columns(df: pd.DataFrame) -> List[str]:
    date_cols = []

    for col in df.columns:
        col_name = str(col).lower()
        if any(keyword in col_name for keyword in ["date", "time", "month", "year"]):
            date_cols.append(str(col))
            continue

        try:
            if df[col].dtype == "object":
                sample = df[col].dropna().astype(str).head(10)
                if len(sample) > 0:
                    parsed = pd.to_datetime(sample, errors="coerce")
                    if parsed.notna().sum() >= max(1, len(sample) // 2):
                        date_cols.append(str(col))
        except Exception:
            pass

    return list(dict.fromkeys(date_cols))[:4]


def build_spreadsheet_kpis(df: pd.DataFrame) -> Dict[str, Any]:
    total_rows = int(len(df))
    total_columns = int(len(df.columns))
    numeric_cols = df.select_dtypes(include="number").columns.tolist()
    date_cols = infer_date_columns(df)

    cards = [
        {"label": "Total Rows", "value": total_rows},
        {"label": "Total Columns", "value": total_columns},
        {"label": "Numeric Columns", "value": len(numeric_cols)},
        {"label": "Missing Cells", "value": int(df.isna().sum().sum())},
    ]

    for col in detect_interesting_numeric_columns(df)[:4]:
        series = pd.to_numeric(df[col], errors="coerce").dropna()
        if len(series) == 0:
            continue
        cards.append({
            "label": f"Sum of {col}",
            "value": round(float(series.sum()), 2)
        })
        if len(cards) >= 8:
            break

    numeric_summaries = []
    for col in detect_interesting_numeric_columns(df):
        series = pd.to_numeric(df[col], errors="coerce").dropna()
        if len(series) == 0:
            continue
        numeric_summaries.append({
            "column": str(col),
            "count": int(series.count()),
            "sum": round(float(series.sum()), 2),
            "average": round(float(series.mean()), 2),
            "min": round(float(series.min()), 2),
            "max": round(float(series.max()), 2),
        })

    column_profiles = []
    for col in df.columns[:12]:
        series = df[col]
        column_profiles.append({
            "column": str(col),
            "dtype": str(series.dtype),
            "non_null": int(series.notna().sum()),
            "nulls": int(series.isna().sum()),
            "unique": int(series.nunique(dropna=True))
        })

    date_summary = []
    for col in date_cols:
        try:
            parsed = pd.to_datetime(df[col], errors="coerce").dropna()
            if len(parsed) > 0:
                date_summary.append({
                    "column": str(col),
                    "min": parsed.min().isoformat(),
                    "max": parsed.max().isoformat()
                })
        except Exception:
            continue

    return {
        "cards": cards,
        "numeric_summaries": numeric_summaries[:6],
        "date_summary": date_summary[:4],
        "column_profiles": column_profiles
    }


def build_csv_or_excel_preview(filename: str, contents: bytes) -> Tuple[str, str, Dict[str, Any]]:
    filename_lower = filename.lower()

    if filename_lower.endswith(".csv"):
        try:
            df = pd.read_csv(io.StringIO(contents.decode("utf-8")))
        except Exception:
            df = pd.read_csv(io.StringIO(contents.decode("latin-1")))

        preview = dataframe_to_preview_rows(df)
        kpis = build_spreadsheet_kpis(df)

        text = df.to_csv(index=False)

        extracted_data = {
            "preview": {
                "kind": "spreadsheet",
                "sheets": [
                    {
                        "sheet_name": "CSV",
                        "columns": preview["columns"],
                        "rows": preview["rows"],
                        "kpis": kpis
                    }
                ],
                "workbook_kpis": kpis
            }
        }

        return text, "csv", extracted_data

    excel_file = pd.ExcelFile(io.BytesIO(contents))
    all_sheet_text = []
    sheets_preview = []
    total_rows_all = 0
    total_columns_max = 0

    for sheet_name in excel_file.sheet_names:
        df = excel_file.parse(sheet_name=sheet_name)
        total_rows_all += len(df)
        total_columns_max = max(total_columns_max, len(df.columns))

        preview = dataframe_to_preview_rows(df)
        kpis = build_spreadsheet_kpis(df)

        sheets_preview.append({
            "sheet_name": str(sheet_name),
            "columns": preview["columns"],
            "rows": preview["rows"],
            "kpis": kpis
        })

        all_sheet_text.append(f"Sheet: {sheet_name}\n")
        all_sheet_text.append(df.to_csv(index=False))

    workbook_cards = [
        {"label": "Sheets", "value": len(excel_file.sheet_names)},
        {"label": "Total Rows", "value": int(total_rows_all)},
        {"label": "Max Columns in Sheet", "value": int(total_columns_max)},
    ]

    extracted_data = {
        "preview": {
            "kind": "spreadsheet",
            "sheets": sheets_preview,
            "workbook_kpis": {
                "cards": workbook_cards
            }
        }
    }

    return "\n\n".join(all_sheet_text), "xlsx", extracted_data


def analyze_image_with_openai(filename: str, contents: bytes) -> Tuple[str, str, Dict[str, Any]]:
    mime_type, _ = mimetypes.guess_type(filename)
    if not mime_type:
        mime_type = "image/png"

    base64_image = base64.b64encode(contents).decode("utf-8")
    data_url = f"data:{mime_type};base64,{base64_image}"

    response = client.responses.create(
        model="gpt-4.1-mini",
        input=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "input_text",
                        "text": """
Analyze this image carefully and return ONLY valid JSON with this shape:
{
  "summary": "short summary",
  "visible_text": "all important visible text",
  "labels": ["object1", "object2"],
  "numbers": ["number1", "number2"],
  "table_like_content": "describe any table if present"
}
"""
                    },
                    {
                        "type": "input_image",
                        "image_url": data_url,
                        "detail": "high"
                    }
                ]
            }
        ]
    )

    raw = response.output_text.strip()
    if raw.startswith("```"):
        raw = raw.replace("```json", "").replace("```", "").strip()

    try:
        parsed = json.loads(raw)
    except Exception:
        parsed = {
            "summary": raw,
            "visible_text": "",
            "labels": [],
            "numbers": [],
            "table_like_content": ""
        }

    summary = parsed.get("summary", "")
    visible_text = parsed.get("visible_text", "")
    labels = parsed.get("labels", [])
    numbers = parsed.get("numbers", [])
    table_like_content = parsed.get("table_like_content", "")

    embedding_text = "\n".join([
        f"Image summary: {summary}",
        f"Visible text: {visible_text}",
        f"Labels: {', '.join(labels) if isinstance(labels, list) else str(labels)}",
        f"Numbers: {', '.join(numbers) if isinstance(numbers, list) else str(numbers)}",
        f"Table-like content: {table_like_content}",
    ]).strip()

    extracted_data = {
        "preview": {
            "kind": "image",
            "summary": summary,
            "visible_text": visible_text,
            "labels": labels,
            "numbers": numbers,
            "table_like_content": table_like_content
        }
    }

    return embedding_text, "image", extracted_data


def extract_file_text(filename: str, contents: bytes):
    filename_lower = filename.lower()
    text = ""
    file_type = ""
    extracted_data: Dict[str, Any] = {}

    if filename_lower.endswith(".csv") or filename_lower.endswith(".xlsx"):
        return build_csv_or_excel_preview(filename, contents)

    elif filename_lower.endswith(".pdf"):
        pdf = fitz.open(stream=contents, filetype="pdf")
        for page in pdf:
            text += page.get_text()
        pdf.close()
        file_type = "pdf"

    elif filename_lower.endswith(".txt"):
        try:
            text = contents.decode("utf-8")
        except Exception:
            text = contents.decode("latin-1")
        file_type = "txt"

    elif filename_lower.endswith(".docx"):
        doc = Document(io.BytesIO(contents))
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        file_type = "docx"

    elif filename_lower.endswith((".png", ".jpg", ".jpeg", ".webp")):
        return analyze_image_with_openai(filename, contents)

    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    if not text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from file")

    return text, file_type, extracted_data


def classify_document(text: str) -> str:
    sample = text[:4000]

    response = client.responses.create(
        model="gpt-4.1-mini",
        input=f"""
Classify this document into exactly one category from this list:
- Spreadsheet
- Image
- Invoice
- Resume
- Loan Document
- Contract
- Research Paper
- Policy
- Other

Return only the category name.

Document text:
{sample}
"""
    )

    return response.output_text.strip()


def extract_document_data(document_type: str, text: str) -> Dict[str, Any]:
    sample = text[:6000]

    prompt_map = {
        "Spreadsheet": """
Extract these fields from the spreadsheet if present:
- dataset_name
- likely_primary_entities
- important_columns
- likely_metrics
- likely_date_columns
- brief_summary

Return valid JSON only.
""",
        "Image": """
Extract these fields from the image if present:
- title_or_subject
- visible_text_summary
- important_objects
- notable_numbers
- brief_summary

Return valid JSON only.
""",
        "Invoice": """
Extract these fields from the invoice if present:
- invoice_number
- invoice_date
- vendor_name
- customer_name
- total_amount
- due_date

Return valid JSON only.
""",
        "Resume": """
Extract these fields from the resume if present:
- full_name
- email
- phone
- skills
- education
- experience_summary

Return valid JSON only.
""",
        "Loan Document": """
Extract these fields from the loan document if present:
- borrower_name
- loan_amount
- interest_rate
- property_address
- loan_term
- due_date

Return valid JSON only.
""",
        "Contract": """
Extract these fields from the contract if present:
- parties
- effective_date
- expiration_date
- contract_value
- governing_law

Return valid JSON only.
""",
        "Research Paper": """
Extract these fields from the research paper if present:
- title
- authors
- abstract_summary
- keywords
- conclusion_summary

Return valid JSON only.
""",
        "Policy": """
Extract these fields from the policy document if present:
- policy_name
- effective_date
- owner
- summary
- key_rules

Return valid JSON only.
""",
        "Other": """
Extract the most important structured fields if present.
Return valid JSON only.
"""
    }

    instruction = prompt_map.get(document_type, prompt_map["Other"])

    response = client.responses.create(
        model="gpt-4.1-mini",
        input=f"""
Return ONLY valid JSON.
Do NOT include markdown fences.
Do NOT include explanations.
Do NOT include text before or after the JSON object.

{instruction}

Document text:
{sample}
"""
    )

    raw = response.output_text.strip()

    if raw.startswith("```"):
        raw = raw.replace("```json", "").replace("```", "").strip()

    try:
        return json.loads(raw)
    except Exception:
        return {"raw_extraction": raw}


def get_relevant_chunks(user_id: str, filename: str, question: str, match_count: int = 3) -> List[Dict[str, Any]]:
    query_embedding = get_embedding(question)

    result = supabase.rpc(
        "match_documents",
        {
            "query_embedding": query_embedding,
            "match_count": match_count,
            "p_user_id": user_id,
            "p_filename": filename,
        },
    ).execute()

    return result.data or []


def build_context_from_chunks(chunks: List[Dict[str, Any]]) -> str:
    return "\n\n".join([row["chunk_text"] for row in chunks if row.get("chunk_text")])


def get_conversation_for_user(conversation_id: str, user_id: str) -> Optional[Dict[str, Any]]:
    convo_result = (
        supabase.table("conversations")
        .select("*")
        .eq("id", conversation_id)
        .eq("user_id", user_id)
        .execute()
    )
    return convo_result.data[0] if convo_result.data else None


def get_document_for_user(user_id: str, filename: str) -> Optional[Dict[str, Any]]:
    result = (
        supabase.table("documents")
        .select("*")
        .eq("user_id", user_id)
        .eq("filename", filename)
        .execute()
    )
    return result.data[0] if result.data else None


def update_conversation_metadata(
    conversation: Dict[str, Any],
    conversation_id: str,
    question: str,
    filename: str,
    now_iso: str
) -> str:
    current_title = conversation.get("title") or "New Chat"
    new_title = current_title

    if current_title == "New Chat":
        shortened = question.strip()
        new_title = shortened[:40] + ("..." if len(shortened) > 40 else "")

    supabase.table("conversations").update({
        "title": new_title,
        "filename": filename,
        "updated_at": now_iso
    }).eq("id", conversation_id).execute()

    return new_title


def insert_usage_event(
    user_id: str,
    conversation_id: Optional[str],
    filename: Optional[str],
    question: str,
    model: str,
    response_time_ms: int,
    now_iso: str
):
    supabase.table("usage_events").insert({
        "user_id": user_id,
        "conversation_id": conversation_id,
        "filename": filename,
        "question": question,
        "model": model,
        "response_time_ms": response_time_ms,
        "created_at": now_iso
    }).execute()


def upload_to_storage(user_id: str, filename: str, contents: bytes, content_type: Optional[str] = None) -> Tuple[str, str]:
    safe_name = filename.replace("/", "_")
    storage_path = f"{user_id}/{uuid.uuid4()}-{safe_name}"

    file_options = {"upsert": "true"}
    if content_type:
        file_options["content-type"] = content_type

    supabase.storage.from_(STORAGE_BUCKET).upload(
        path=storage_path,
        file=contents,
        file_options=file_options
    )

    public_url = supabase.storage.from_(STORAGE_BUCKET).get_public_url(storage_path)
    return storage_path, public_url


def build_excel_bytes_from_document(document: Dict[str, Any]) -> bytes:
    extracted_data = document.get("extracted_data") or {}
    preview = extracted_data.get("preview") or {}
    structured_fields = extracted_data.get("structured_fields") or {}

    output = io.BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        file_type = document.get("file_type")

        if file_type in ["csv", "xlsx"]:
            sheets = preview.get("sheets") or []
            if sheets:
                for idx, sheet in enumerate(sheets):
                    sheet_name = str(sheet.get("sheet_name") or f"Sheet{idx + 1}")[:31]
                    rows = sheet.get("rows") or []
                    df = pd.DataFrame(rows)
                    if df.empty:
                        df = pd.DataFrame(columns=sheet.get("columns") or [])
                    df.to_excel(writer, sheet_name=sheet_name, index=False)

                    kpis = (sheet.get("kpis") or {}).get("cards") or []
                    if kpis:
                        kpi_df = pd.DataFrame(kpis)
                        kpi_sheet_name = f"{sheet_name[:24]}_KPIs"[:31]
                        kpi_df.to_excel(writer, sheet_name=kpi_sheet_name, index=False)
            else:
                pd.DataFrame([{"Message": "No spreadsheet preview found"}]).to_excel(
                    writer, sheet_name="Report", index=False
                )
        else:
            summary_rows = []

            if preview.get("summary"):
                summary_rows.append({"Field": "Summary", "Value": preview.get("summary")})
            if preview.get("visible_text"):
                summary_rows.append({"Field": "Visible Text", "Value": preview.get("visible_text")})
            if preview.get("labels"):
                summary_rows.append({"Field": "Labels", "Value": ", ".join(preview.get("labels"))})
            if preview.get("numbers"):
                summary_rows.append({"Field": "Numbers", "Value": ", ".join(preview.get("numbers"))})
            if preview.get("table_like_content"):
                summary_rows.append({"Field": "Table-like Content", "Value": preview.get("table_like_content")})

            if structured_fields:
                for key, value in structured_fields.items():
                    if isinstance(value, list):
                        value = ", ".join([str(v) for v in value])
                    elif isinstance(value, dict):
                        value = json.dumps(value)
                    summary_rows.append({"Field": str(key), "Value": str(value)})

            summary_df = pd.DataFrame(summary_rows or [{"Field": "Info", "Value": "No structured report available"}])
            summary_df.to_excel(writer, sheet_name="Report", index=False)

    output.seek(0)
    return output.getvalue()


def build_pdf_bytes_from_document(document: Dict[str, Any]) -> bytes:
    extracted_data = document.get("extracted_data") or {}
    preview = extracted_data.get("preview") or {}
    structured_fields = extracted_data.get("structured_fields") or {}

    output = io.BytesIO()
    pdf = canvas.Canvas(output, pagesize=letter)
    _, height = letter
    y = height - 50

    def write_line(text: str, font_size: int = 11, gap: int = 16):
        nonlocal y
        if y < 60:
            pdf.showPage()
            y = height - 50
        pdf.setFont("Helvetica", font_size)
        pdf.drawString(50, y, text[:110])
        y -= gap

    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(50, y, "AI Document Report")
    y -= 24

    write_line(f"Filename: {document.get('filename', '')}")
    write_line(f"File Type: {document.get('file_type', '')}")
    write_line(f"Document Type: {document.get('document_type', '')}")
    write_line("")

    file_type = document.get("file_type")

    if file_type in ["csv", "xlsx"]:
        workbook_kpis = (preview.get("workbook_kpis") or {}).get("cards") or []
        if workbook_kpis:
            pdf.setFont("Helvetica-Bold", 13)
            pdf.drawString(50, y, "Workbook KPIs")
            y -= 20
            for card in workbook_kpis:
                write_line(f"{card.get('label')}: {card.get('value')}")

        sheets = preview.get("sheets") or []
        for sheet in sheets[:3]:
            if y < 120:
                pdf.showPage()
                y = height - 50
            pdf.setFont("Helvetica-Bold", 13)
            pdf.drawString(50, y, f"Sheet: {sheet.get('sheet_name')}")
            y -= 20

            for card in ((sheet.get("kpis") or {}).get("cards") or [])[:6]:
                write_line(f"{card.get('label')}: {card.get('value')}")
    else:
        if preview.get("summary"):
            pdf.setFont("Helvetica-Bold", 13)
            pdf.drawString(50, y, "Summary")
            y -= 20
            write_line(str(preview.get("summary")))

        if preview.get("visible_text"):
            pdf.setFont("Helvetica-Bold", 13)
            pdf.drawString(50, y, "Visible Text")
            y -= 20
            for line in str(preview.get("visible_text")).split("\n")[:12]:
                write_line(line)

        if preview.get("labels"):
            pdf.setFont("Helvetica-Bold", 13)
            pdf.drawString(50, y, "Labels")
            y -= 20
            write_line(", ".join(preview.get("labels")))

    if structured_fields:
        if y < 120:
            pdf.showPage()
            y = height - 50
        pdf.setFont("Helvetica-Bold", 13)
        pdf.drawString(50, y, "Structured Fields")
        y -= 20
        for key, value in list(structured_fields.items())[:15]:
            if isinstance(value, list):
                value = ", ".join([str(v) for v in value])
            elif isinstance(value, dict):
                value = json.dumps(value)
            write_line(f"{key}: {value}")

    pdf.save()
    output.seek(0)
    return output.getvalue()


def get_history_for_conversation(conversation_id: str, user_id: str) -> List[Dict[str, Any]]:
    result = (
        supabase.table("chat_history")
        .select("*")
        .eq("conversation_id", conversation_id)
        .eq("user_id", user_id)
        .order("created_at", desc=False)
        .execute()
    )
    return result.data or []


def add_snapshot_to_doc(doc: Document, snapshot_base64: Optional[str]) -> None:
    if not snapshot_base64:
        return

    try:
        base64_data = snapshot_base64.split(",")[-1]
        image_bytes = base64.b64decode(base64_data)
        image_stream = io.BytesIO(image_bytes)

        doc.add_heading("Visual Snapshot", level=2)
        doc.add_picture(image_stream, width=Inches(6.5))
    except Exception as e:
        print("SNAPSHOT INSERT ERROR:", str(e))


def build_docx_bytes_from_document(
    document: Dict[str, Any],
    conversation_messages: Optional[List[Dict[str, Any]]] = None,
    snapshot_base64: Optional[str] = None
) -> bytes:
    extracted_data = document.get("extracted_data") or {}
    preview = extracted_data.get("preview") or {}
    structured_fields = extracted_data.get("structured_fields") or {}

    doc = Document()
    doc.add_heading("AI Document Report", level=1)
    doc.add_paragraph(f"Filename: {document.get('filename', '')}")
    doc.add_paragraph(f"File Type: {document.get('file_type', '')}")
    doc.add_paragraph(f"Document Type: {document.get('document_type', '')}")

    file_type = document.get("file_type")

    if file_type in ["csv", "xlsx"]:
        workbook_kpis = (preview.get("workbook_kpis") or {}).get("cards") or []
        if workbook_kpis:
            doc.add_heading("Workbook KPIs", level=2)
            for card in workbook_kpis:
                doc.add_paragraph(f"{card.get('label')}: {card.get('value')}")

        sheets = preview.get("sheets") or []
        for sheet in sheets[:3]:
            doc.add_heading(f"Sheet: {sheet.get('sheet_name')}", level=2)

            cards = ((sheet.get("kpis") or {}).get("cards") or [])[:8]
            if cards:
                for card in cards:
                    doc.add_paragraph(f"{card.get('label')}: {card.get('value')}")

            rows = (sheet.get("rows") or [])[:8]
            cols = sheet.get("columns") or []

            if rows and cols:
                table = doc.add_table(rows=1, cols=len(cols))
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(cols):
                    hdr_cells[i].text = str(col)

                for row in rows:
                    row_cells = table.add_row().cells
                    for i, col in enumerate(cols):
                        val = row.get(col, "")
                        row_cells[i].text = "" if val is None else str(val)
    else:
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(str(preview.get("summary") or "No summary available"))

        if preview.get("visible_text"):
            doc.add_heading("Visible Text", level=2)
            doc.add_paragraph(str(preview.get("visible_text")))

        if preview.get("labels"):
            doc.add_heading("Labels", level=2)
            doc.add_paragraph(", ".join(preview.get("labels")))

        if preview.get("numbers"):
            doc.add_heading("Numbers", level=2)
            doc.add_paragraph(", ".join(preview.get("numbers")))

        if preview.get("table_like_content"):
            doc.add_heading("Table-like Content", level=2)
            doc.add_paragraph(str(preview.get("table_like_content")))

    if structured_fields:
        doc.add_heading("Structured Fields", level=2)
        for key, value in structured_fields.items():
            if isinstance(value, list):
                value = ", ".join([str(v) for v in value])
            elif isinstance(value, dict):
                value = json.dumps(value)
            doc.add_paragraph(f"{key}: {value}")

    if conversation_messages:
        doc.add_heading("Conversation History", level=2)
        for index, item in enumerate(conversation_messages, start=1):
            q = item.get("question", "")
            a = item.get("answer", "")
            created_at = item.get("created_at", "")

            doc.add_paragraph(f"Q{index}: {q}")
            doc.add_paragraph(f"A{index}: {a}")

            if created_at:
                doc.add_paragraph(f"Created At: {created_at}")

    add_snapshot_to_doc(doc, snapshot_base64)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


@app.get("/")
def root():
    return {"status": "ok"}


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/loan-risk")
async def loan_risk_prediction(data: LoanApplication):
    if not ML_API_URL:
        raise HTTPException(
            status_code=500,
            detail="ML_API_URL is not configured"
        )

    payload = data.dict()
    last_error = None

    for attempt in range(4):
        try:
            print(f"Waking ML API attempt {attempt + 1}/4")

            # Wake up Render service first
            try:
                requests.get(ML_API_URL, timeout=60)
            except Exception as wake_error:
                print("Wake-up request failed:", str(wake_error))

            print(f"Calling ML predict attempt {attempt + 1}/4")

            response = requests.post(
                f"{ML_API_URL}/predict",
                json=payload,
                timeout=120
            )

            response.raise_for_status()
            return response.json()

        except requests.exceptions.RequestException as e:
            last_error = str(e)
            print(f"ML API attempt {attempt + 1} failed: {last_error}")

            if attempt < 3:
                wait_time = 10 * (attempt + 1)
                print(f"Waiting {wait_time} seconds before retry...")
                time.sleep(wait_time)

    raise HTTPException(
        status_code=502,
        detail=f"ML API request failed after wake-up retries: {last_error}"
    )

@app.get("/documents/{user_id}")
def get_documents(user_id: str):
    try:
        result = (
            supabase.table("documents")
            .select("*")
            .eq("user_id", user_id)
            .order("uploaded_at", desc=True)
            .execute()
        )
        return result.data
    except Exception as e:
        print("DOCUMENTS ERROR:", str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/documents/{user_id}/{filename}")
def delete_document(user_id: str, filename: str):
    try:
        docs_result = (
            supabase.table("documents")
            .select("storage_path")
            .eq("user_id", user_id)
            .eq("filename", filename)
            .execute()
        )

        for row in (docs_result.data or []):
            storage_path = row.get("storage_path")
            if storage_path:
                try:
                    supabase.storage.from_(STORAGE_BUCKET).remove([storage_path])
                except Exception:
                    pass

        convo_result = (
            supabase.table("conversations")
            .select("id")
            .eq("user_id", user_id)
            .eq("filename", filename)
            .execute()
        )

        convo_ids = [row["id"] for row in (convo_result.data or [])]
        for convo_id in convo_ids:
            supabase.table("chat_history").delete().eq("conversation_id", convo_id).execute()

        supabase.table("conversations").delete().eq("user_id", user_id).eq("filename", filename).execute()
        supabase.table("document_chunks").delete().eq("user_id", user_id).eq("filename", filename).execute()
        supabase.table("document_chunks_v2").delete().eq("user_id", user_id).eq("filename", filename).execute()
        supabase.table("documents").delete().eq("user_id", user_id).eq("filename", filename).execute()

        return {"message": "Document and related chats deleted"}
    except Exception as e:
        print("DELETE DOCUMENT ERROR:", str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/conversations/{user_id}")
def get_conversations(user_id: str, filename: Optional[str] = None, search: Optional[str] = None):
    try:
        query = (
            supabase.table("conversations")
            .select("*")
            .eq("user_id", user_id)
        )

        if filename:
            query = query.eq("filename", filename)

        if search:
            query = query.or_(f"title.ilike.%{search}%,filename.ilike.%{search}%")

        result = query.order("updated_at", desc=True).execute()
        return result.data
    except Exception as e:
        print("CONVERSATIONS ERROR:", str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/conversations")
def create_conversation(request: CreateConversationRequest):
    try:
        now_iso = utc_now_iso()
        result = supabase.table("conversations").insert({
            "user_id": request.user_id,
            "title": request.title or "New Chat",
            "filename": request.filename,
            "created_at": now_iso,
            "updated_at": now_iso,
            "is_public": False,
            "share_token": None,
        }).execute()

        if not result.data:
            raise HTTPException(status_code=500, detail="Failed to create conversation")

        return result.data[0]
    except Exception as e:
        print("CREATE CONVERSATION ERROR:", str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.patch("/conversations/{conversation_id}")
def update_conversation(conversation_id: str, request: UpdateConversationRequest):
    try:
        update_data = {
            "updated_at": utc_now_iso()
        }

        if request.title is not None:
            update_data["title"] = request.title

        if request.filename is not None:
            update_data["filename"] = request.filename

        result = (
            supabase.table("conversations")
            .update(update_data)
            .eq("id", conversation_id)
            .execute()
        )

        return result.data[0] if result.data else {}
    except Exception as e:
        print("UPDATE CONVERSATION ERROR:", str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/conversations/{conversation_id}")
def delete_conversation(conversation_id: str):
    try:
        supabase.table("chat_history").delete().eq("conversation_id", conversation_id).execute()
        supabase.table("conversations").delete().eq("id", conversation_id).execute()
        return {"message": "Conversation deleted"}
    except Exception as e:
        print("DELETE CONVERSATION ERROR:", str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/conversations/{conversation_id}/share")
def share_conversation(conversation_id: str):
    try:
        token = str(uuid.uuid4())
        result = (
            supabase.table("conversations")
            .update({
                "is_public": True,
                "share_token": token,
                "updated_at": utc_now_iso()
            })
            .eq("id", conversation_id)
            .execute()
        )

        if not result.data:
            raise HTTPException(status_code=404, detail="Conversation not found")

        return {
            "share_token": token,
            "share_path": f"?share={token}"
        }
    except Exception as e:
        print("SHARE CONVERSATION ERROR:", str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/shared/{share_token}")
def get_shared_conversation(share_token: str):
    try:
        convo_result = (
            supabase.table("conversations")
            .select("*")
            .eq("share_token", share_token)
            .eq("is_public", True)
            .execute()
        )

        if not convo_result.data:
            raise HTTPException(status_code=404, detail="Shared conversation not found")

        conversation = convo_result.data[0]

        history_result = (
            supabase.table("chat_history")
            .select("*")
            .eq("conversation_id", conversation["id"])
            .order("created_at", desc=False)
            .execute()
        )

        return {
            "conversation": conversation,
            "messages": history_result.data or []
        }
    except HTTPException:
        raise
    except Exception as e:
        print("GET SHARED CONVERSATION ERROR:", str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/history/{conversation_id}")
def get_chat_history(conversation_id: str, user_id: Optional[str] = None):
    try:
        query = (
            supabase.table("chat_history")
            .select("*")
            .eq("conversation_id", conversation_id)
        )

        if user_id:
            query = query.eq("user_id", user_id)

        result = query.order("created_at", desc=False).execute()
        return result.data
    except Exception as e:
        print("HISTORY ERROR:", str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/upload")
async def upload_file(file: UploadFile = File(...), user_id: str = Form(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file")

    try:
        contents = await file.read()

        text, file_type, preview_data = extract_file_text(file.filename, contents)
        chunks = chunk_text(text)

        document_type = classify_document(text)
        structured_data = extract_document_data(document_type, text)

        storage_path, public_url = upload_to_storage(
            user_id=user_id,
            filename=file.filename,
            contents=contents,
            content_type=file.content_type
        )

        supabase.table("document_chunks").delete().eq("user_id", user_id).eq("filename", file.filename).execute()
        supabase.table("document_chunks_v2").delete().eq("user_id", user_id).eq("filename", file.filename).execute()
        supabase.table("documents").delete().eq("user_id", user_id).eq("filename", file.filename).execute()

        rows_v2 = []
        now_iso = utc_now_iso()

        for idx, chunk in enumerate(chunks):
            emb = get_embedding(chunk)
            rows_v2.append({
                "user_id": user_id,
                "filename": file.filename,
                "chunk_index": idx,
                "chunk_text": chunk,
                "embedding": emb,
                "created_at": now_iso
            })

        batch_size = 50
        for i in range(0, len(rows_v2), batch_size):
            supabase.table("document_chunks_v2").insert(rows_v2[i:i + batch_size]).execute()

        extracted_data = {
            **preview_data,
            "structured_fields": structured_data
        }

        supabase.table("documents").insert({
            "user_id": user_id,
            "filename": file.filename,
            "file_type": file_type,
            "uploaded_at": now_iso,
            "document_type": document_type,
            "extracted_data": extracted_data,
            "storage_path": storage_path,
            "public_url": public_url
        }).execute()

        return {
            "message": "Upload successful",
            "filename": file.filename,
            "chunks": len(chunks),
            "saved_to_supabase": True,
            "vector_table": "document_chunks_v2",
            "document_type": document_type,
            "file_type": file_type,
            "extracted_data": extracted_data,
            "public_url": public_url
        }

    except HTTPException:
        raise
    except Exception as e:
        print("UPLOAD ERROR:", str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/ask")
def ask_ai(request: AskRequest):
    start_time = time.time()

    try:
        if not request.conversation_id:
            return {"answer": "No active conversation selected.", "saved_to_supabase": False}

        conversation = get_conversation_for_user(request.conversation_id, request.user_id)

        if not conversation:
            return {"answer": "Conversation not found.", "saved_to_supabase": False}

        selected_filename = request.filename or conversation.get("filename")

        if not selected_filename:
            return {"answer": "No document selected for this conversation.", "saved_to_supabase": False}

        matched_chunks = get_relevant_chunks(
            request.user_id,
            selected_filename,
            request.question,
            match_count=3
        )

        if not matched_chunks:
            return {
                "answer": f"No stored chunks found for '{selected_filename}'. Please upload it again.",
                "saved_to_supabase": False
            }

        context = build_context_from_chunks(matched_chunks)

        response = client.responses.create(
            model="gpt-4.1-mini",
            input=f"""
You are inside an AI document platform.

Answer using only this context.

Important:
- If the user asks to export or create a PDF, Excel, DOCX/Word, or snapshot, do NOT say you cannot do it.
- Instead, briefly say that the export can be done using the platform action.
- Keep the answer short and helpful.

Context:
{context}

Question:
{request.question}
"""
        )

        final_answer = response.output_text
        now_iso = utc_now_iso()

        supabase.table("chat_history").insert({
            "user_id": request.user_id,
            "conversation_id": request.conversation_id,
            "question": request.question,
            "answer": final_answer,
            "filename": selected_filename,
            "created_at": now_iso
        }).execute()

        new_title = update_conversation_metadata(
            conversation=conversation,
            conversation_id=request.conversation_id,
            question=request.question,
            filename=selected_filename,
            now_iso=now_iso
        )

        response_time_ms = int((time.time() - start_time) * 1000)

        insert_usage_event(
            user_id=request.user_id,
            conversation_id=request.conversation_id,
            filename=selected_filename,
            question=request.question,
            model="gpt-4.1-mini",
            response_time_ms=response_time_ms,
            now_iso=now_iso
        )

        return {
            "answer": final_answer,
            "saved_to_supabase": True,
            "conversation_id": request.conversation_id,
            "title": new_title,
            "response_time_ms": response_time_ms,
            "retrieved_count": len(matched_chunks)
        }

    except Exception as e:
        print("ASK ERROR:", str(e))
        return {"answer": f"Error: {str(e)}", "saved_to_supabase": False}


@app.post("/ask-stream")
def ask_ai_stream(request: AskRequest):
    start_time = time.time()

    try:
        if not request.conversation_id:
            raise HTTPException(status_code=400, detail="No active conversation selected.")

        conversation = get_conversation_for_user(request.conversation_id, request.user_id)

        if not conversation:
            raise HTTPException(status_code=404, detail="Conversation not found.")

        selected_filename = request.filename or conversation.get("filename")

        if not selected_filename:
            raise HTTPException(status_code=400, detail="No document selected for this conversation.")

        matched_chunks = get_relevant_chunks(
            request.user_id,
            selected_filename,
            request.question,
            match_count=3
        )

        if not matched_chunks:
            raise HTTPException(status_code=404, detail=f"No stored chunks found for '{selected_filename}'. Please upload it again.")

        context = build_context_from_chunks(matched_chunks)

        def generate():
            collected: List[str] = []

            stream = client.responses.create(
                model="gpt-4.1-mini",
                input=f"""
You are inside an AI document platform.

Answer using only this context.

Important:
- If the user asks to export or create a PDF, Excel, DOCX/Word, or snapshot, do NOT say you cannot do it.
- Instead, briefly say that the export can be done using the platform action.
- Keep the answer short and helpful.

Context:
{context}

Question:
{request.question}
""",
                stream=True
            )

            for event in stream:
                if event.type == "response.output_text.delta":
                    delta = event.delta
                    collected.append(delta)
                    yield f"data: {json.dumps({'token': delta})}\n\n"

            final_answer = "".join(collected)
            now_iso = utc_now_iso()

            supabase.table("chat_history").insert({
                "user_id": request.user_id,
                "conversation_id": request.conversation_id,
                "question": request.question,
                "answer": final_answer,
                "filename": selected_filename,
                "created_at": now_iso
            }).execute()

            new_title = update_conversation_metadata(
                conversation=conversation,
                conversation_id=request.conversation_id,
                question=request.question,
                filename=selected_filename,
                now_iso=now_iso
            )

            response_time_ms = int((time.time() - start_time) * 1000)

            insert_usage_event(
                user_id=request.user_id,
                conversation_id=request.conversation_id,
                filename=selected_filename,
                question=request.question,
                model="gpt-4.1-mini",
                response_time_ms=response_time_ms,
                now_iso=now_iso
            )

            payload = {
                "done": True,
                "conversation_id": request.conversation_id,
                "title": new_title,
                "response_time_ms": response_time_ms,
                "retrieved_count": len(matched_chunks)
            }

            yield f"data: {json.dumps(payload)}\n\n"

        return StreamingResponse(generate(), media_type="text/event-stream")

    except HTTPException:
        raise
    except Exception as e:
        print("ASK STREAM ERROR:", str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/export-excel-report")
def export_excel_report(request: ExportReportRequest):
    try:
        document = get_document_for_user(request.user_id, request.filename)
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        excel_bytes = build_excel_bytes_from_document(document)
        base_name = request.filename.rsplit(".", 1)[0] if "." in request.filename else request.filename

        return StreamingResponse(
            io.BytesIO(excel_bytes),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={base_name}_report.xlsx"}
        )
    except HTTPException:
        raise
    except Exception as e:
        print("EXPORT EXCEL ERROR:", str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/export-pdf-report")
def export_pdf_report(request: ExportReportRequest):
    try:
        document = get_document_for_user(request.user_id, request.filename)
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        pdf_bytes = build_pdf_bytes_from_document(document)
        base_name = request.filename.rsplit(".", 1)[0] if "." in request.filename else request.filename

        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={base_name}_report.pdf"}
        )
    except HTTPException:
        raise
    except Exception as e:
        print("EXPORT PDF ERROR:", str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/export-docx-report")
def export_docx_report(request: ExportReportRequest):
    try:
        document = get_document_for_user(request.user_id, request.filename)
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        docx_bytes = build_docx_bytes_from_document(document)
        base_name = request.filename.rsplit(".", 1)[0] if "." in request.filename else request.filename

        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={base_name}_report.docx"}
        )
    except HTTPException:
        raise
    except Exception as e:
        print("EXPORT DOCX ERROR:", str(e))
        raise HTTPException(status_code=500, detail=str(e))